from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse, HttpResponse
from django.shortcuts import get_object_or_404
from docx.shared import Pt
from django.contrib import messages

from docx import Document
from .models import Certificate 
import io
import json

def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']  # We'll ignore the password here
        print(f"Username: {username}, Password: {password}")

        # Check if the user already exists in the database
        try:
            user = User.objects.get(username=username)
        except User.DoesNotExist:
            # If the user does not exist, create a new one
            user = User.objects.create_user(username=username, password='dummy_password')
            user.save()

        # Automatically log in the user without checking password
        login(request, user)
        print("Login successful")
        return redirect('admin_panel')

    return render(request, 'login.html')

@login_required(login_url='login')  # This protects the admin panel
def admin_panel_view(request):
    completed_interns_count = Certificate.objects.filter(status="completed").count()
    pending_interns_count = Certificate.objects.filter(status="pending").count()
    print(f"Completed Interns: {completed_interns_count}, Pending Interns: {pending_interns_count}")


    # Pass the counts to the template
    context = {
        'completed_interns_count': completed_interns_count,
        'pending_interns_count': pending_interns_count
    }
   
    # Check if the user is authenticated
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login if not authenticated
    return render(request, 'admin_panel.html',context)
def new_intern_view(request):
    return render(request, 'new_intern.html')  
def generate_certificate(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)

            # Select the appropriate document based on nameTitle
            if data['nameTitle'] == 'Mr.':
                template_file = "dummy_certificate_male.docx"
            else:  # Assuming Ms. or Mrs. falls into this category
                template_file = "dummy_certificate_female.docx"

            # Open the selected document
            with open(template_file, "r") as doc_file:
                doc = Document(template_file)
            
            # Replace placeholders with actual data
            for paragraph in doc.paragraphs:
                paragraph.text = paragraph.text.replace('{{Name}}', data['name'])
                paragraph.text = paragraph.text.replace('{{Name Title}}', data['nameTitle'])
                paragraph.text = paragraph.text.replace('{{Degree}}', data['degree'])
                paragraph.text = paragraph.text.replace('{{College Name}}', data['collegeName'])
                paragraph.text = paragraph.text.replace('{{Department}}',data['Department'])
                paragraph.text = paragraph.text.replace('{{From Date}}', data['fromDate'])
                paragraph.text = paragraph.text.replace('{{To Date}}', data['toDate'])
                paragraph.text = paragraph.text.replace('{{Location}}', data['location'])
                paragraph.text = paragraph.text.replace('{{Guide Name}}', data['guideName'])
                paragraph.text = paragraph.text.replace('{{Guide Title}}', data['guideTitle'])
                paragraph.text = paragraph.text.replace('{{Guide Designation}}', data['guideDesignation'])
                paragraph.text = paragraph.text.replace('{{Guide Location}}', data['guideLocation'])
                paragraph.text = paragraph.text.replace('{{Guide CO Title}}', data['guideCOTitle'])
                paragraph.text = paragraph.text.replace('{{Guide CO Name}}', data['guideCOName'])
                paragraph.text = paragraph.text.replace('{{Guide CO Designation}}', data['guideCODesignation'])
                paragraph.text = paragraph.text.replace('{{Guide CO Location}}', data['guideCOLocation'])
                paragraph.text = paragraph.text.replace('{{Date}}', data['date'])

                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        run.bold = False  # Ensure the text is unbol
                        run.font.size = Pt(12)  # Example of setting font size
                        run.font.name = 'Times New Roman'
            # Save the modified document to an in-memory buffer
            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            # Return the generated certificate as a Word file
            response = HttpResponse(file_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=certificate.docx'
            return response

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Invalid request method'}, status=405)

def save_form_data(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            
            # Access data directly from the JSON payload
            name = data.get('name')
            name_title = data.get('nameTitle')
            print(name_title)
            Department=data.get('Department')
            guide_title = data.get('guideTitle')
            print(guide_title)
            guide_co_title = data.get('guideCOTitle')
            print(guide_co_title)
            degree = data.get('degree')
            college_name = data.get('collegeName')
            from_date = data.get('fromDate')
            to_date = data.get('toDate')
            location = data.get('location')
            guide_name = data.get('guideName')
            guide_designation = data.get('guideDesignation')
            guide_location = data.get('guideLocation')
            guide_co_name = data.get('guideCOName')
            guide_co_designation = data.get('guideCODesignation')
            guide_co_location = data.get('guideCOLocation')
            date = data.get('date')
            
            # Create the Certificate instance and save it
            certificate = Certificate(  
                name=name,
                name_title=name_title,
                degree=degree,  
                college_name=college_name, 
                Department=Department, 
                from_date=from_date,  
                to_date=to_date,  
                location=location,  
                guide_name=guide_name,
                guide_title=guide_title,  
                guide_designation=guide_designation,  
                guide_location=guide_location,  
                guide_co_name=guide_co_name,
                guide_co_title=guide_co_title,  
                guide_co_designation=guide_co_designation,  
                guide_co_location=guide_co_location,  
                date=date  
            )

            # Save form data into the database
            certificate.save()  
            return JsonResponse({'message': 'Form data saved successfully!'}, status=200)

        except Exception as e:
            return JsonResponse({'error': str(e)}, status=400)

    return JsonResponse({'error': 'Invalid request method'}, status=405)
def database(request):
    certificates = Certificate.objects.all()  # Fetching all Certificate records
    return render(request, 'database.html', {'certificates': certificates}) 
@login_required
def edit_certificate(request, id):
    certificate = get_object_or_404(Certificate, id=id)
    
    if request.method == 'POST':
        certificate.name = request.POST['name']
        certificate.degree = request.POST['degree']
        certificate.college_name = request.POST['college_name']
        certificate.from_date = request.POST['from_date']
        certificate.to_date = request.POST['to_date']
        certificate.Department = request.POST['Department']
        certificate.location = request.POST['location']
        certificate.guide_name = request.POST['guide_name']
        certificate.guide_designation = request.POST['guide_designation']
        certificate.guide_location = request.POST['guide_location']
        certificate.guide_co_name = request.POST['guide_co_name']
        certificate.guide_co_designation = request.POST['guide_co_designation']
        certificate.guide_co_location = request.POST['guide_co_location']

        
        certificate.save()
        messages.success(request, "Certificate updated successfully!")
        return redirect('database')
    
    return render(request, 'edit_certificate.html', {'certificate': certificate})


@login_required
def delete_certificate(request, id):
    certificate = get_object_or_404(Certificate, id=id)
    certificate.delete()
    messages.success(request, "Certificate deleted  successfully!")
    return redirect('database')
def Existing_Intern(request):
    certificates = Certificate.objects.all()  # Fetching all Certificate records
    return render(request, 'Existing_intern.html', {'certificates': certificates})
def update_status(request, certificate_id):
    
    if request.method == 'POST':
        # Fetch the certificate first
        certificate = get_object_or_404(Certificate, id=certificate_id)
        
        # Now safely set its status
        certificate.status = request.POST.get('status')
        certificate.save()
        return redirect('Existing_Intern')
    context = {
        'certificate_id': certificate_id,  # Pass the certificate_id to the template
    }
    return render(request, 'Existing_Intern',context)
def Generate_Completion_Certificate(request, certificate_id):
    # Get the certificate data based on the provided certif):
    context = {
        'certificate_id': certificate_id,  # Pass the certificate_id to the template
    }
    return render(request, 'Generate_Completion_Certificate.html', context) 
def generate_completion_certificate(request, certificate_id):
    # Get the certificate data ba
    # paragraph.text = paragraph.text.replace('{{Department}}',data['Department'])sed on the provided certificate_id
    certificate = get_object_or_404(Certificate, id=certificate_id)

    if request.method == 'POST':
        # Retrieve the project title from the request
        project_title = request.POST.get('project_title')

        # Select the appropriate document based on name_title
        if certificate.name_title == 'Mr.':
            template_file = "dummy_certificate_completion_male.docx"
        else:  # Assuming Ms. or Mrs. falls into this category
            template_file = "dummy_certificate_completion_female.docx"

        # Load the selected dummy document
        doc = Document(template_file)  # Use the selected template file

        # Replace placeholders with actual data
        for paragraph in doc.paragraphs:
            paragraph.text = paragraph.text.replace('{{Name}}', certificate.name)
            paragraph.text = paragraph.text.replace('{{Name Title}}', certificate.name_title)
            paragraph.text = paragraph.text.replace('{{Degree}}', certificate.degree)
            paragraph.text = paragraph.text.replace('{{College Name}}', certificate.college_name)
            paragraph.text = paragraph.text.replace('{{From Date}}', str(certificate.from_date))
            paragraph.text = paragraph.text.replace('{{To Date}}', str(certificate.to_date))
            paragraph.text = paragraph.text.replace('{{Department}}',certificate.Department)
            paragraph.text = paragraph.text.replace('{{Location}}', certificate.location)
            paragraph.text = paragraph.text.replace('{{Guide Name}}', certificate.guide_name)
            paragraph.text = paragraph.text.replace('{{Guide Title}}', certificate.guide_title)
            paragraph.text = paragraph.text.replace('{{Guide Designation}}', certificate.guide_designation)
            paragraph.text = paragraph.text.replace('{{Guide Location}}', certificate.guide_location)
            paragraph.text = paragraph.text.replace('{{Guide CO Name}}', certificate.guide_co_name)
            paragraph.text = paragraph.text.replace('{{Guide CO Title}}', certificate.guide_co_title)
            paragraph.text = paragraph.text.replace('{{Guide CO Designation}}', certificate.guide_co_designation)
            paragraph.text = paragraph.text.replace('{{Guide CO Location}}', certificate.guide_co_location)
            paragraph.text = paragraph.text.replace('{{Date}}', str(certificate.date))
            paragraph.text = paragraph.text.replace('{{Project Title}}', project_title)

        # Save the modified document to a BytesIO stream
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={certificate.name}_completion_certificate.docx'

        # Write the document to the response
        doc.save(response)

        return response

    return HttpResponse(status=405)  # Return a 405 Method Not Allowed if not POSTeturn a 405 Method Not Allowed if not POST
@login_required(login_url='login')
def admin_dashboard(request):
    # Fetch counts of completed and pending interns
    completed_interns_count = Certificate.objects.filter(status="Completed").count()
    pending_interns_count = Certificate.objects.filter(status="Pending").count()
    print(f"Completed Interns: {completed_interns_count}, Pending Interns: {pending_interns_count}")


    # Pass the counts to the template
    context = {
        'completed_interns_count': completed_interns_count,
        'pending_interns_count': pending_interns_count
    }
    return render(request, 'admin_panel.html', context)